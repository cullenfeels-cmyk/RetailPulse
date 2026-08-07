import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette Constants
    HEX_PRIMARY = "1B365D"      # Deep Navy
    HEX_SECONDARY = "4A777A"    # Slate Teal
    HEX_DARK = "333333"         # Charcoal Body Text
    HEX_LIGHT_BG = "F4F6F9"     # Light Table Shading

    # Helper function for cell background color
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Helper function for custom styled headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(27, 54, 93)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(74, 119, 122)
        return p

    # Helper function for standard body paragraphs
    def add_body_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.bold = True
            r_bold.font.color.rgb = RGBColor(51, 51, 51)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    # --- DOCUMENT TITLE / COVER HEADER ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run("RetailPulse AI Sales Intelligence Platform")
    t_run.font.name = 'Arial'
    t_run.font.size = Pt(22)
    t_run.bold = True
    t_run.font.color.rgb = RGBColor(27, 54, 93)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_p.add_run("Comprehensive Internship Project Report & Documentation")
    s_run.font.name = 'Arial'
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = RGBColor(74, 119, 122)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 1. EXECUTIVE SUMMARY ---
    add_custom_heading("1. Executive Summary", level=1)
    add_body_p(
        "RetailPulse AI is a comprehensive, production-grade retail analytics platform engineered to empower businesses with actionable sales insights, demand forecasts, customer segmentation, churn prediction, and inventory health optimization. Built using Python, Pandas, Plotly, and Streamlit, the platform processes millions of historical transactional data points to deliver a multi-page interactive web application.",
        bold_prefix="Overview: "
    )

    # --- 2. PROJECT OBJECTIVES ---
    add_custom_heading("2. Project Objectives", level=1)
    add_body_p("Predicting future sales volumes using time-series modeling to ensure optimized supply chain planning.", bold_prefix="• Demand Forecasting: ")
    add_body_p("Gauging customer behavior through Recency, Frequency, and Monetary (RFM) segmentation models.", bold_prefix="• Customer Intelligence: ")
    add_body_p("Identifying high-risk accounts based on inactivity thresholds to prevent customer churn.", bold_prefix="• Churn Prevention: ")
    add_body_p("Balancing stock levels using ABC inventory categorization and economic order quantity principles.", bold_prefix="• Inventory Optimization: ")

    # --- 3. TECHNICAL METHODOLOGY ---
    add_custom_heading("3. Technical Methodology", level=1)
    add_body_p("Ingested historical retail datasets using Pandas, handling missing values, filtering invalid prices, and removing canceled transaction records.", bold_prefix="Phase 1 - Data Ingestion & ETL: ")
    add_body_p("Engineered custom behavioral features including customer recency, order frequency, total monetary spend, and monthly time-series aggregations.", bold_prefix="Phase 2 - Feature Engineering: ")
    add_body_p("Developed modular analytical components across 15 specialized views integrated seamlessly into a unified Streamlit application dashboard.", bold_prefix="Phase 3 - Application Deployment: ")

    # --- 4. KEY PERFORMANCE METRICS TABLE ---
    add_custom_heading("4. Key Performance Summary Table", level=1)
    add_body_p("The table below outlines the core metrics extracted from the global transactional dataset:")

    table_data = [
        ["Metric Category", "Indicator Value", "Analytical Focus"],
        ["Total Revenue", "10.27M", "Overall sales financial output"],
        ["Total Orders", "21K", "Transactional volume throughput"],
        ["Total Customers", "4,313", "Unique active buyer accounts"],
        ["Global Reach", "40 Countries", "International market presence"],
        ["Average Order Value", "490.28", "Mean basket value per transaction"]
    ]

    table = doc.add_table(rows=len(table_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            
            if i == 0:
                set_cell_background(cell, HEX_PRIMARY)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if i % 2 == 1:
                    set_cell_background(cell, HEX_LIGHT_BG)
                run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- 5. TECHNOLOGY STACK ---
    add_custom_heading("5. Technology Stack", level=1)
    add_body_p("Python 3.12 for core data processing and algorithmic computation.", bold_prefix="• Programming Language: ")
    add_body_p("Streamlit framework with custom CSS styling and native multi-page routing.", bold_prefix="• Dashboard Framework: ")
    add_body_p("Plotly Express for dynamic, interactive data visualizations and charts.", bold_prefix="• Visualization Engine: ")
    add_body_p("Git and GitHub for robust version control and deployment integration.", bold_prefix="• Deployment & Version Control: ")

    # --- 6. CONCLUSION ---
    add_custom_heading("6. Conclusion", level=1)
    add_body_p(
        "RetailPulse AI successfully bridges the gap between raw data science modeling and practical business intelligence. By combining robust multi-page analytics, real-time filtering, and professional UI layout, the platform delivers a scalable solution for modern retail environments.",
        bold_prefix="Summary: "
    )

    # Save document
    filename = "RetailPulse_Project_Report.docx"
    doc.save(filename)
    print(f"Successfully generated and saved: {filename}")

if __name__ == "__main__":
    create_report()