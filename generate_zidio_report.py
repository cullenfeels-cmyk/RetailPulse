import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_full_zidio_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styling Palette
    HEX_PRIMARY = "1B365D"      # Deep Navy
    HEX_SECONDARY = "4A777A"    # Slate Teal
    HEX_LIGHT_BG = "F4F6F9"     # Light Shading

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.bold = True
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(27, 54, 93)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(74, 119, 122)
        return p

    def add_body_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.bold = True
            r_bold.font.color.rgb = RGBColor(51, 51, 51)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    # --- COVER TITLE ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run("RetailPulse - AI-Powered Customer Analytics & Demand Forecasting")
    t_run.font.name = 'Arial'
    t_run.font.size = Pt(18)
    t_run.bold = True
    t_run.font.color.rgb = RGBColor(27, 54, 93)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_p.add_run("Zidio Internship Comprehensive Project Report & Documentation")
    s_run.font.name = 'Arial'
    s_run.font.size = Pt(12)
    s_run.font.color.rgb = RGBColor(74, 119, 122)

    # --- 1. PROJECT OVERVIEW ---
    add_custom_heading("1. Project Overview", level=1)
    add_body_p(
        "RetailPulse is an AI-powered retail analytics platform designed to help businesses understand customer behavior, forecast product demand, and optimize inventory management using Machine Learning and MLOps practices. The project focuses on improving business intelligence through predictive analytics and real-time visual dashboards[cite: 3].",
        bold_prefix="Executive Summary: "
    )
    
    add_custom_heading("Vision & Objectives:", level=2)
    add_body_p("Predict future product demand using AI forecasting models[cite: 3].", bold_prefix="• ")
    add_body_p("Analyze customer purchasing behavior[cite: 3].", bold_prefix="• ")
    add_body_p("Identify churn-prone customers[cite: 3].", bold_prefix="• ")
    add_body_p("Improve inventory planning and reduce overstocking[cite: 3].", bold_prefix="• ")

    add_custom_heading("Business Value Delivered:", level=2)
    add_body_p("Better forecasting accuracy[cite: 3].", bold_prefix="• ")
    add_body_p("Reduced inventory wastage[cite: 3].", bold_prefix="• ")
    add_body_p("Improved customer retention[cite: 3].", bold_prefix="• ")
    add_body_p("Data-driven decision making[cite: 3].", bold_prefix="• ")

    # --- 2. ARCHITECTURE OVERVIEW ---
    add_custom_heading("2. System Architecture & MLOps Workflow", level=1)
    add_body_p(
        "RetailPulse follows an end-to-end AI and MLOps pipeline structure: Raw Sales Data → Data Cleaning → Feature Engineering → Customer Segmentation & Churn Prediction → Forecasting Models (Prophet + LSTM + Hybrid) → Inventory Optimization → Streamlit Dashboard → MLflow Tracking & Monitoring[cite: 3]."
    )
    add_body_p("The system is modular, scalable, and incorporates industry-standard experiment tracking, model monitoring, retraining readiness, and workflow automation[cite: 3].", bold_prefix="Architecture Design: ")

    # --- 3. EXECUTION TIMELINE ---
    add_custom_heading("3. Execution Timeline", level=1)
    add_body_p("Requirement gathering, dataset collection, preprocessing, and Exploratory Data Analysis (EDA)[cite: 3].", bold_prefix="Week 1 (24 April - 30 April 2026): ")
    add_body_p("Feature engineering, customer segmentation using RFM analysis, and churn prediction implementation[cite: 3].", bold_prefix="Week 2 (1 May - 7 May 2026): ")
    add_body_p("Demand forecasting using Prophet and LSTM, hybrid model integration, and inventory optimization model building[cite: 3].", bold_prefix="Week 3 (8 May - 15 May 2026): ")
    add_body_p("Streamlit dashboard development, Docker containerization, deployment testing, and final report preparation[cite: 3].", bold_prefix="Week 4 (16 May - 24 May 2026): ")

    # --- 4. TECHNICAL HIGHLIGHTS & CHALLENGES ---
    add_custom_heading("4. Technical Highlights & Solutions", level=1)
    add_body_p("Customer Segmentation (K-Means & RFM), Churn Prediction (Classification), Demand Forecasting (Prophet & LSTM), and Hybrid model integration[cite: 3].", bold_prefix="Modeling Techniques: ")
    add_body_p("Customer purchase frequency, recency, monetary value, sales trend extraction, and time-series transformations[cite: 3].", bold_prefix="Feature Engineering: ")
    add_body_p("MLflow for experiment tracking, Optuna for hyperparameter tuning, data drift detection, and Airflow orchestration[cite: 3].", bold_prefix="MLOps Practices: ")
    add_body_p("Handled inconsistent sales data via robust preprocessing pipelines and resolved deployment configurations using containerized Docker optimization[cite: 3].", bold_prefix="Challenges & Solutions: ")

    # --- 5. KEY FEATURES TABLE ---
    add_custom_heading("5. Key Features Summary Table", level=1)
    
    features_data = [
        ["ID", "Feature Name", "Description", "Acceptance Criteria"],
        ["F1", "Customer Segmentation", "Groups customers based on behavior", "Customers categorized successfully[cite: 3]"],
        ["F2", "Churn Prediction", "Predicts at-risk customers", "Accurate churn insights generated[cite: 3]"],
        ["F3", "Demand Forecasting", "Predicts future demand", "Reliable forecast trends shown[cite: 3]"],
        ["F4", "Hybrid Forecasting", "Combines multiple models", "Improved prediction quality[cite: 3]"],
        ["F5", "Inventory Optimization", "Optimizes stock levels", "Reduced overstocking risk[cite: 3]"],
        ["F6", "Dashboard Analytics", "Interactive visualization", "Real-time insights accessible[cite: 3]"],
        ["F7", "Drift Detection", "Detects data changes", "Alerts for model degradation[cite: 3]"],
        ["F8", "MLflow Tracking", "Tracks experiments", "Performance history maintained[cite: 3]"]
    ]

    table1 = doc.add_table(rows=len(features_data), cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(features_data):
        row = table1.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            if i == 0:
                set_cell_background(cell, HEX_PRIMARY)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if i % 2 == 1:
                    set_cell_background(cell, HEX_LIGHT_BG)
                run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 6. TECHNOLOGY STACK TABLE ---
    add_custom_heading("6. Technology Stack", level=1)

    tech_data = [
        ["Category", "Technology", "Rationale / Alternatives"],
        ["Programming", "Python", "Primary language for ML and analytics[cite: 3]"],
        ["Frontend", "Streamlit", "Interactive dashboard development[cite: 3]"],
        ["Machine Learning", "Scikit-learn, TensorFlow", "Model training and prediction[cite: 3]"],
        ["Forecasting", "Prophet, LSTM", "Time-series forecasting[cite: 3]"],
        ["MLOps", "MLflow, Optuna, Airflow", "Experiment tracking and orchestration[cite: 3]"],
        ["Containerization", "Docker", "Portable deployment[cite: 3]"],
        ["Version Control", "Git & GitHub", "Collaboration and code management[cite: 3]"]
    ]

    table2 = doc.add_table(rows=len(tech_data), cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(tech_data):
        row = table2.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            if i == 0:
                set_cell_background(cell, HEX_PRIMARY)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if i % 2 == 1:
                    set_cell_background(cell, HEX_LIGHT_BG)
                run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 7. DEPLOYMENT & CONCLUSION ---
    add_custom_heading("7. Deployment & Personal Reflection", level=1)
    add_body_p("Hosted via Streamlit Cloud, containerized with Docker, and tracked using GitHub version control[cite: 3].", bold_prefix="Deployment Platforms: ")
    add_body_p(
        "This project provided practical exposure to real-world Machine Learning and MLOps concepts[cite: 3]. During development, we learned how to design scalable data pipelines, build predictive models, and deploy applications professionally[cite: 3].",
        bold_prefix="Conclusion: "
    )

    filename = "RetailPulse_Complete_Zidio_Report.docx"
    doc.save(filename)
    print(f"Successfully generated complete report: {filename}")

if __name__ == "__main__":
    create_full_zidio_report()